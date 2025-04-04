import csv
from docx import Document
from docx.shared import Pt
from datetime import datetime

def generate_word_doc(account_number, activities):
    doc = Document()
    doc.add_heading(f"Account Statement for Account Number: {account_number}", level=1)
    doc.add_paragraph("--------------------------------------------------")

    for timestamp, activity in activities:
        p = doc.add_paragraph()
        run = p.add_run(f"{timestamp}: {activity}")
        run.font.size = Pt(12)

    doc.save(f"account_statement_{account_number}.docx")

def account_statement(account_number):
    activities = []
    try:
        with open('activity_log.csv', newline='') as logfile:
            reader = csv.DictReader(logfile)
            for row in reader:
                if int(row['Account Number']) == account_number:
                    activities.append((row['Timestamp'], row['Activity']))
    except FileNotFoundError:
        print("Activity log file not found.")
        return
    
    if activities:
        print(f"Account Statement for Account Number: {account_number}")
        print("--------------------------------------------------")
        for timestamp, activity in activities:
            print(f"{timestamp}: {activity}")
        generate_word_doc(account_number, activities)
        print("Word document generated successfully.")
    else:
        print("No activities found for this account.")


# test_docx.py
# from docx import Document

# doc = Document()
# doc.add_heading('Test Document', level=1)
# doc.add_paragraph('This is a test document to verify the installation of python-docx.')

# doc.save('test_document.docx')
# print("Document created successfully.")
